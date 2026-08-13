import os
import docx
import pandas as pd
from spire.doc import Document, FileFormat
from src.schema import ExtractedDocument, DocumentChunk

class TextParser:
    @staticmethod
    def parse(file_path: str, file_name: str) -> ExtractedDocument:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
        chunk = DocumentChunk(
            content=text.strip(),
            metadata={"source": file_name, "type": "txt"}
        )
        return ExtractedDocument(
            file_name=file_name,
            file_type="txt",
            chunks=[chunk],
            full_text=text.strip()
        )


class DocxParser:
    @staticmethod
    def parse(file_path: str, file_name: str) -> ExtractedDocument:
        doc = docx.Document(file_path)
        chunks = []
        full_text_list = []

        # 1. Extract Paragraphs
        for idx, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:  # Ignore empty lines
                chunks.append(DocumentChunk(
                    content=text,
                    metadata={"source": file_name, "paragraph_index": idx, "type": "docx"}
                ))
                full_text_list.append(text)

        # 2. Extract Embedded Tables
        for t_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_rows.append(" | ".join(row_data))
            
            formatted_table = "\n".join(table_rows)
            if formatted_table:
                content = f"--- WORD TABLE {t_idx + 1} ---\n" + formatted_table
                chunks.append(DocumentChunk(
                    content=content,
                    metadata={"source": file_name, "table_index": t_idx, "type": "docx_table"}
                ))
                full_text_list.append(formatted_table)

        full_text = "\n\n".join(full_text_list)
        return ExtractedDocument(
            file_name=file_name,
            file_type="docx",
            chunks=chunks,
            full_text=full_text
        )


class DocParser:
    @staticmethod
    def parse(file_path: str, file_name: str) -> ExtractedDocument:
        temp_docx_path = f"{file_path}.temp.docx"
        try:
            # 1. Convert legacy .doc to .docx on the fly
            doc = Document()
            doc.LoadFromFile(file_path)
            doc.SaveToFile(temp_docx_path, FileFormat.Docx)
            doc.Close()

            # 2. Delegate to DocxParser
            extracted = DocxParser.parse(temp_docx_path, file_name)
            
            # 3. Return ExtractedDocument with file_type='doc'
            return ExtractedDocument(
                file_name=file_name,
                file_type="doc",
                chunks=extracted.chunks,
                full_text=extracted.full_text
            )
        finally:
            # Clean up temporary converted docx file
            if os.path.exists(temp_docx_path):
                try:
                    os.remove(temp_docx_path)
                except OSError:
                    pass


class ExcelParser:
    @staticmethod
    def parse(file_path: str, file_name: str) -> ExtractedDocument:
        chunks = []
        full_text_list = []

        # Handle CSV or Excel
        is_csv = os.path.splitext(file_name)[1].lower() == ".csv"
        if is_csv:
            sheets_dict = {"CSV_Data": pd.read_csv(file_path)}
        else:
            sheets_dict = pd.read_excel(file_path, sheet_name=None)  # Reads all sheets

        for sheet_name, df in sheets_dict.items():
            # Clean completely empty rows and columns
            df = df.dropna(how="all").dropna(axis=1, how="all")
            
            if df.empty:
                continue

            # Convert dataframe to Markdown format so LLM understands table columns/rows
            markdown_table = df.to_markdown(index=False)
            content = f"### Sheet: {sheet_name}\n\n{markdown_table}"
            
            chunks.append(DocumentChunk(
                content=content,
                metadata={
                    "source": file_name,
                    "sheet_name": sheet_name,
                    "row_count": len(df),
                    "columns": list(df.columns),
                    "type": "excel"
                }
            ))
            full_text_list.append(content)

        full_text = "\n\n".join(full_text_list)
        return ExtractedDocument(
            file_name=file_name,
            file_type="csv" if is_csv else "xlsx",
            chunks=chunks,
            full_text=full_text
        )
